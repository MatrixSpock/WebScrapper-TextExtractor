from PIL import Image
import pytesseract
from docx import Document
import os

def extract_text_from_images(image_directory):
    """
    Extract text from all images in the specified directory and create a Word document
    
    Args:
        image_directory (str): Path to directory containing images
    """
    # Create a new Word document
    doc = Document()
    doc.add_heading('Extracted Text from Images', 0)
    
    # Process each image in the directory
    for filename in os.listdir(image_directory):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            try:
                # Full path to image
                image_path = os.path.join(image_directory, filename)
                
                # Open image
                image = Image.open(image_path)
                
                # Extract text from image
                text = pytesseract.image_to_string(image)
                
                # Add image filename as heading
                doc.add_heading(f'Text from {filename}', level=1)
                
                # Add extracted text to document
                doc.add_paragraph(text)
                
                # Add a page break after each image's content
                doc.add_page_break()
                
                print(f"Processed {filename}")
                
            except Exception as e:
                print(f"Error processing {filename}: {str(e)}")
    
    # Save the document
    output_path = os.path.join(image_directory, 'extracted_text.docx')
    doc.save(output_path)
    print(f"\nDocument saved as: {output_path}")

def main():
    # Get directory path from user
    directory = input("Enter the path to your image directory: ")
    
    if os.path.isdir(directory):
        extract_text_from_images(directory)
    else:
        print("Invalid directory path!")

if __name__ == "__main__":
    main()