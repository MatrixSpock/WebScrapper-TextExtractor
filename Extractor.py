from PIL import Image
import pytesseract
from docx import Document
import os

# Ensure the folder exists for output
os.makedirs('./Extractions', exist_ok=True)

# Load the uploaded image
image_path = r'.\screenshots\android.png'  # Use raw strings
image = Image.open(image_path)

# Perform OCR to extract text
extracted_text = pytesseract.image_to_string(image)

# Process the text to extract questions (assuming questions end with '?')
lines = extracted_text.splitlines()
questions = [line.strip() for line in lines if '?' in line]

# Create a .docx file and add the questions
doc = Document()
doc.add_heading('Extracted Questions', level=1)

for i, question in enumerate(questions, start=1):
    doc.add_paragraph(f"{i}. {question}")

# Save the .docx file
output_path = r'.\Extractions\Extracted_Questions2.docx'  # Use raw strings
doc.save(output_path)

print(f"Document saved to: {output_path}")
